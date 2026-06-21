from pathlib import Path

from docx import Document as DocxDocument

from app.loaders.base import BaseLoader
from app.rag.schema import RawDocument


class DocLoader(BaseLoader):
    supported_extensions = (".docx", ".doc")

    def load(self, file_path: Path) -> list[RawDocument]:
        if file_path.suffix.lower() == ".doc":
            raise NotImplementedError("Legacy .doc 文件建议先转换为 .docx 后再解析。")

        document = DocxDocument(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        return [RawDocument(content=text, source_file=file_path.name, file_type="docx")]
